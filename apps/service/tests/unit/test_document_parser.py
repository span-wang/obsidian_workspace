from __future__ import annotations

from hashlib import sha256
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

import workers.document_parser as document_parser
from workers.document_parser import DocumentParseError, parse_document, parse_items


def _write_pdf(path: Path, page_texts: list[list[str]]) -> None:
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids [{' '.join(f'{index + 3} 0 R' for index in range(len(page_texts)))}] /Count {len(page_texts)} >>".encode(),
    ]
    page_objects: list[bytes] = []
    content_objects: list[bytes] = []
    for index, lines in enumerate(page_texts):
        content_number = index + 3 + len(page_texts)
        page_objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 {len(page_texts) * 2 + 3} 0 R >> >> "
                f"/Contents {content_number} 0 R >>"
            ).encode()
        )
        commands = ["BT", "/F1 12 Tf", "72 720 Td"]
        for line_index, line in enumerate(lines):
            if line_index:
                commands.append("0 -20 Td")
            commands.append(f"({_escape_pdf_literal(line)}) Tj")
        commands.append("ET")
        stream = "\n".join(commands).encode()
        content_objects.append(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")
    objects.extend(page_objects)
    objects.extend(content_objects)
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode())
        output.extend(body)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    output.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode())
    path.write_bytes(output)


def _escape_pdf_literal(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def test_parse_pdf_keeps_page_locators_and_detects_question_answer_relationship(tmp_path: Path) -> None:
    source = tmp_path / "lesson.pdf"
    _write_pdf(
        source,
        [
            ["Chapter One", "1. List point"],
            ["Question: What is evidence?", "Answer: A source-linked result."],
        ],
    )

    evidence = parse_document(source, "pdf")

    assert evidence.document_kind == "pdf"
    assert evidence.raw_extraction["pages"][0]["page"] == 1
    assert evidence.raw_extraction["pages"][1]["page"] == 2
    assert any(unit.kind == "heading" and unit.locator.page == 1 for unit in evidence.units)
    assert any(unit.kind == "list-item" and unit.locator.page == 1 for unit in evidence.units)
    assert any(unit.kind == "question-answer" and unit.locator.page == 2 for unit in evidence.units)
    assert evidence.issues == ()


def test_parse_docx_keeps_heading_list_table_and_equivalent_locations(tmp_path: Path) -> None:
    source = tmp_path / "lesson.docx"
    document = Document()
    document.add_heading("Unit One", level=1)
    document.add_paragraph("Vocabulary", style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Question"
    table.cell(0, 1).text = "Answer"
    document.add_paragraph("Question: Where is the source?")
    document.add_paragraph("Answer: In the vault.")
    document.save(source)

    evidence = parse_document(source, "docx")

    assert evidence.document_kind == "docx"
    assert any(unit.kind == "heading" and unit.locator.docx_location == "paragraph:1" for unit in evidence.units)
    assert any(unit.kind == "list-item" and unit.locator.docx_location == "paragraph:2" for unit in evidence.units)
    assert any(unit.kind == "table-cell" and unit.locator.docx_location == "table:1/row:1/cell:1" for unit in evidence.units)
    assert any(unit.kind == "question-answer" for unit in evidence.units)
    assert all(unit.locator.page is None for unit in evidence.units)


def test_parse_docx_without_normal_style_keeps_unstyled_paragraph(tmp_path: Path) -> None:
    source = tmp_path / "unstyled.docx"
    document = Document()
    document.add_paragraph("Unstyled evidence")
    normal = document.styles["Normal"]
    document.styles.element.remove(normal._element)
    document.save(source)

    evidence = parse_document(source, "docx")

    assert any(
        unit.kind == "paragraph"
        and unit.text == "Unstyled evidence"
        and unit.locator.docx_location == "paragraph:1"
        for unit in evidence.units
    )


def test_parse_empty_pdf_creates_a_locatable_required_check_issue(tmp_path: Path) -> None:
    source = tmp_path / "empty.pdf"
    _write_pdf(source, [[]])

    evidence = parse_document(source, "pdf")

    assert evidence.units == ()
    assert evidence.issues[0].code == "empty-page"
    assert evidence.issues[0].locator.page == 1
    assert evidence.issues[0].severity == "required-check"


def test_parse_zero_page_pdf_uses_a_document_region_not_a_docx_locator(tmp_path: Path) -> None:
    source = tmp_path / "zero-pages.pdf"
    writer = PdfWriter()
    with source.open("wb") as output:
        writer.write(output)

    evidence = parse_document(source, "pdf")

    locator = evidence.issues[0].locator
    assert evidence.issues[0].code == "missing-pages"
    assert locator.page is None
    assert locator.docx_location is None
    assert locator.region == "document"


def test_parse_rejects_unknown_or_unreadable_documents(tmp_path: Path) -> None:
    source = tmp_path / "broken.pdf"
    source.write_bytes(b"not a PDF")

    with pytest.raises(DocumentParseError):
        parse_document(source, "pdf")


def test_parse_worker_events_keep_raw_text_off_the_event_log(tmp_path: Path) -> None:
    source = tmp_path / "lesson.pdf"
    _write_pdf(source, [["Chapter One"]])

    events = list(
        parse_items(({"item_id": 7, "path": str(source), "document_kind": "pdf"},))
    )

    assert events[0] == {"type": "parse-started"}
    assert events[1]["type"] == "parse-item"
    assert events[1]["item_id"] == 7
    assert events[1]["content_sha256"] == sha256(source.read_bytes()).hexdigest()
    assert events[1]["evidence"]["raw_extraction"]["pages"][0]["text"] == "Chapter One"
    assert events[-1] == {"type": "parse-completed"}


def test_parse_worker_hides_source_paths_when_the_file_disappears(tmp_path: Path) -> None:
    source = tmp_path / "removed.pdf"

    events = list(parse_items(({"item_id": 7, "path": str(source), "document_kind": "pdf"},)))

    assert events[1] == {
        "type": "parse-failed-item",
        "item_id": 7,
        "reason": "The source file is no longer available for local parsing.",
        "locator_summary": "document",
    }
    assert str(source) not in events[1]["reason"]


def test_parse_worker_uses_one_snapshot_when_the_source_changes_during_parsing(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "lesson.pdf"
    _write_pdf(source, [["Original chapter"]])
    original_bytes = source.read_bytes()
    parse_pdf = document_parser._parse_pdf

    def replace_source_after_snapshot(source_bytes: bytes, should_cancel):
        _write_pdf(source, [["Replaced chapter"]])
        return parse_pdf(source_bytes, should_cancel)

    monkeypatch.setattr(document_parser, "_parse_pdf", replace_source_after_snapshot)
    events = list(parse_items(({"item_id": 7, "path": str(source), "document_kind": "pdf"},)))

    assert events[1]["content_sha256"] == sha256(original_bytes).hexdigest()
    assert events[1]["evidence"]["raw_extraction"]["pages"][0]["text"] == "Original chapter"


def test_parse_worker_cancels_inside_a_document(tmp_path: Path) -> None:
    source = tmp_path / "lesson.pdf"
    _write_pdf(source, [["Chapter One"], ["Chapter Two"]])
    cancellation_checks = iter((False, False, True))

    events = list(
        parse_items(
            ({"item_id": 7, "path": str(source), "document_kind": "pdf"},),
            should_cancel=lambda: next(cancellation_checks),
        )
    )

    assert events == [{"type": "parse-started"}, {"type": "parse-cancelled"}]


def test_docx_preserves_body_order_heading_levels_and_atomic_question_answers(tmp_path: Path) -> None:
    source = tmp_path / "ordered.docx"
    document = Document()
    document.add_heading("Chapter", level=1)
    document.add_paragraph("Before table")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table content"
    document.add_heading("Scope", level=2)
    document.add_paragraph("Question: Why?")
    document.add_paragraph("Answer: Because.")
    document.add_paragraph("Afterword")
    document.save(source)

    evidence = parse_document(source, "docx")

    assert [(unit.kind, unit.text) for unit in evidence.units] == [
        ("heading", "Chapter"),
        ("paragraph", "Before table"),
        ("table-cell", "Table content"),
        ("heading-2", "Scope"),
        ("question-answer", "Question: Why?\nAnswer: Because."),
        ("paragraph", "Afterword"),
    ]


def test_pdf_question_answer_is_one_atomic_unit_without_a_duplicate_answer(tmp_path: Path) -> None:
    source = tmp_path / "question.pdf"
    _write_pdf(source, [["Question: Why?", "Answer: Because."]])

    evidence = parse_document(source, "pdf")

    assert [(unit.kind, unit.text) for unit in evidence.units] == [
        ("question-answer", "Question: Why?\nAnswer: Because."),
    ]
