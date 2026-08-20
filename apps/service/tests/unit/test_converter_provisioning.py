from __future__ import annotations

import json
import os
from hashlib import sha256
from pathlib import Path
from types import SimpleNamespace

import api.main as api_main
from api.runtime import RuntimeState
from docx import Document as WordDocument
from domain.derived_notes import render_document_graph
from domain.evidence import ArtifactRef, DocumentGraph, PdfRegionLocator
from workers.converters.launcher import (
    _converter_environment,
    _adapt_graph,
    _artifacts,
    _docling_blocks,
    _mineru_blocks,
    _pandoc_blocks,
)
from workers.converters.profiles import require_profile
from workers.converters.provisioning import (
    ProvisionedProfiles,
    default_converter_root,
    load_provisioned_profiles,
)
from workers.document_parser import preflight_document


def test_loader_verifies_only_a_controlled_manifest_and_keeps_approval_gates_closed(
    tmp_path: Path,
) -> None:
    root = tmp_path / "ObsidianPlatform" / "converters"
    executable = root / "pandoc-3.10" / "Pandoc" / "pandoc.exe"
    config = root / "profiles" / "pandoc.json"
    model = root / "models" / "pandoc-assets.json"
    executable.parent.mkdir(parents=True)
    config.parent.mkdir(parents=True)
    model.parent.mkdir(parents=True)
    executable.write_bytes(b"fixed pandoc executable")
    config.write_text('{"offline":true}', encoding="utf-8")
    model.write_bytes(b"fixed supporting asset")
    _write_manifest(
        root,
        {
            "profile_id": "pandoc-3.10-local",
            "engine": "pandoc",
            "engine_version": "3.10",
            "executable": "pandoc-3.10/Pandoc/pandoc.exe",
            "executable_sha256": _digest(executable),
            "config": "profiles/pandoc.json",
            "config_sha256": _digest(config),
            "models": [{"path": "models/pandoc-assets.json", "sha256": _digest(model)}],
            "resource_limits": {"wall_clock_seconds": 60, "memory_mb": 4096},
            # These untrusted manifest claims must not enable a local engine.
            "release_approved": True,
            "network_denied": True,
        },
    )

    loaded = load_provisioned_profiles(root)
    profile = loaded.profile_for("pandoc")

    assert profile is not None
    assert Path(profile.executable_path or "").resolve() == executable.resolve()
    assert profile.model_hashes == (_digest(model),)
    assert profile.release_approved is False
    assert profile.network_denied is False
    assert profile.isolation_boundary == "local-process"
    assert require_profile(profile, "pandoc").reason_code == "release-approval-missing"


def test_loader_rejects_path_escape_and_changed_hashes(tmp_path: Path) -> None:
    root = tmp_path / "converters"
    root.mkdir()
    outside = tmp_path / "outside.exe"
    outside.write_bytes(b"outside")
    _write_manifest(
        root,
        {
            "profile_id": "escaped",
            "engine": "pandoc",
            "engine_version": "3.10",
            "executable": "../outside.exe",
            "executable_sha256": _digest(outside),
            "config": "missing.json",
            "config_sha256": "0" * 64,
            "models": [],
            "resource_limits": {"wall_clock_seconds": 60},
        },
    )

    escaped = load_provisioned_profiles(root)

    assert escaped.profile_for("pandoc") is None
    assert escaped.unavailable_reasons["pandoc"] == "profile-integrity-invalid"

    executable = root / "pandoc.exe"
    config = root / "pandoc.json"
    executable.write_bytes(b"old binary")
    config.write_bytes(b"fixed config")
    _write_manifest(
        root,
        {
            "profile_id": "changed",
            "engine": "pandoc",
            "engine_version": "3.10",
            "executable": "pandoc.exe",
            "executable_sha256": "0" * 64,
            "config": "pandoc.json",
            "config_sha256": _digest(config),
            "models": [],
            "resource_limits": {"wall_clock_seconds": 60},
        },
    )

    changed = load_provisioned_profiles(root)

    assert changed.profile_for("pandoc") is None
    assert changed.unavailable_reasons["pandoc"] == "profile-integrity-invalid"


def test_default_root_uses_only_localappdata_converter_location(monkeypatch) -> None:
    monkeypatch.setenv("LOCALAPPDATA", r"C:\\Users\\example\\AppData\\Local")
    monkeypatch.setenv("PATH", r"C:\\somewhere\\with\\pandoc")

    assert default_converter_root() == Path(r"C:\\Users\\example\\AppData\\Local") / "ObsidianPlatform" / "converters"


def test_composition_root_injects_one_private_store_and_unavailable_profile_map(
    tmp_path: Path, monkeypatch
) -> None:
    root = tmp_path / "converters"
    monkeypatch.setattr(
        api_main,
        "load_provisioned_profiles",
        lambda: ProvisionedProfiles(
            root,
            {},
            {
                "mineru": "profile-missing",
                "pandoc": "profile-missing",
                "docling": "profile-missing",
                "paddleocr-vl": "profile-missing",
            },
        ),
    )
    runtime = RuntimeState(data_directory=tmp_path / "app-data", sqlite_version="3.45.1")

    app = api_main.create_app(runtime=runtime)
    service = app.state.import_task_service
    worker = service.worker

    assert service.converter_profile == {}
    assert service.artifact_store is worker._artifact_store
    assert service.artifact_store.root == runtime.data_directory / "conversion-artifacts"


def test_loader_requires_a_hash_bound_local_approval_record(tmp_path: Path) -> None:
    root = tmp_path / "converters"
    executable = root / "pandoc.exe"
    config = root / "pandoc.json"
    root.mkdir()
    executable.write_bytes(b"pandoc")
    config.write_bytes(b"offline")
    profile = {
        "profile_id": "pandoc-local",
        "engine": "pandoc",
        "engine_version": "3.10",
        "executable": "pandoc.exe",
        "executable_sha256": _digest(executable),
        "config": "pandoc.json",
        "config_sha256": _digest(config),
        "models": [],
        "resource_limits": {"wall_clock_seconds": 60},
    }
    _write_manifest(root, profile)
    (root / "converter-release-approval.json").write_text(
        json.dumps(
            {
                "schema_version": 1,
                "approved_profiles": [
                    {
                        "engine": "pandoc",
                        "profile_id": "pandoc-local",
                        "executable_sha256": _digest(executable),
                        "config_hash": _digest(config),
                        "model_hashes": [],
                        "license_disposition": "local-use",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    profile = load_provisioned_profiles(root).profile_for("pandoc")

    assert profile is not None
    assert profile.release_approved is True
    assert require_profile(profile, "pandoc").allowed is True


def test_loader_accepts_only_an_approved_native_paddleocr_vl_profile(tmp_path: Path) -> None:
    root = tmp_path / "converters"
    executable = root / "paddleocr" / "paddleocr.exe"
    config = root / "profiles" / "paddleocr-vl-1.6.yaml"
    model = root / "models" / "PaddleOCR-VL-1.6-0.9B.bin"
    executable.parent.mkdir(parents=True)
    config.parent.mkdir(parents=True)
    model.parent.mkdir(parents=True)
    executable.write_bytes(b"paddleocr executable")
    config.write_text("pipeline_name: PaddleOCR-VL-1.6\n", encoding="utf-8")
    model.write_bytes(b"local model")
    profile_entry = {
        "profile_id": "paddleocr-vl-1.6-local",
        "engine": "paddleocr-vl",
        "engine_version": "3.7.0",
        "executable": "paddleocr/paddleocr.exe",
        "executable_sha256": _digest(executable),
        "config": "profiles/paddleocr-vl-1.6.yaml",
        "config_sha256": _digest(config),
        "models": [{"path": "models/PaddleOCR-VL-1.6-0.9B.bin", "sha256": _digest(model)}],
        "resource_limits": {"wall_clock_seconds": 60, "memory_mb": 4096},
        "backends": ["native"],
    }
    _write_manifest(root, profile_entry)
    (root / "converter-release-approval.json").write_text(
        json.dumps(
            {
                "schema_version": 1,
                "approved_profiles": [
                    {
                        "engine": "paddleocr-vl",
                        "profile_id": profile_entry["profile_id"],
                        "executable_sha256": profile_entry["executable_sha256"],
                        "config_hash": profile_entry["config_sha256"],
                        "model_hashes": [_digest(model)],
                        "license_disposition": "local-use",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    profile = load_provisioned_profiles(root).profile_for("paddleocr-vl")

    assert profile is not None
    assert profile.release_approved is True
    assert profile.supports_backend("native") is True
    assert require_profile(profile, "paddleocr-vl").allowed is True


def test_converter_environment_keeps_paddleocr_vl_offline(tmp_path: Path) -> None:
    profile = SimpleNamespace(executable_path=r"C:\\converter\\paddleocr.exe", config_path=None)

    environment = _converter_environment(profile, tmp_path)

    assert environment["HF_HUB_OFFLINE"] == "1"
    assert environment["TRANSFORMERS_OFFLINE"] == "1"
    assert environment["PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK"] == "True"
    assert environment["PADDLE_PDX_CACHE_HOME"] == str(tmp_path / "paddle-cache")
    assert "HTTP_PROXY" not in environment
    assert "HTTPS_PROXY" not in environment


def test_converter_environment_uses_the_preprovisioned_paddleocr_cache(tmp_path: Path) -> None:
    model = tmp_path / "paddleocr-vl-1.6" / "official_models" / "PaddleOCR-VL-1.6"
    model.mkdir(parents=True)
    profile = SimpleNamespace(
        engine="paddleocr-vl",
        executable_path=r"C:\\converter\\paddleocr.exe",
        config_path=None,
        model_paths=(str(model),),
    )

    environment = _converter_environment(profile, tmp_path / "attempt")

    assert environment["PADDLE_PDX_CACHE_HOME"] == str(model.parent.parent)


def test_converter_environment_adds_only_paddle_runtime_dll_directories(tmp_path: Path) -> None:
    executable = tmp_path / "runtime" / "Scripts" / "paddleocr.exe"
    paddle_libs = tmp_path / "runtime" / "Lib" / "site-packages" / "paddle" / "libs"
    nvidia_bin = tmp_path / "runtime" / "Lib" / "site-packages" / "nvidia" / "cudnn" / "bin"
    executable.parent.mkdir(parents=True)
    paddle_libs.mkdir(parents=True)
    nvidia_bin.mkdir(parents=True)
    profile = SimpleNamespace(
        engine="paddleocr-vl",
        executable_path=str(executable),
        config_path=None,
        model_paths=(),
    )

    environment = _converter_environment(profile, tmp_path / "attempt")

    path_entries = environment["PATH"].split(os.pathsep)
    assert path_entries == [str(executable.parent), str(nvidia_bin), str(paddle_libs)]


def test_pandoc_adapter_skips_layout_only_docx_paragraphs_and_supports_quotes(
    tmp_path: Path,
) -> None:
    source = tmp_path / "quoted.docx"
    document = WordDocument()
    document.add_paragraph("Opening evidence")
    document.add_paragraph()
    document.add_paragraph("First quoted evidence")
    document.add_paragraph("Second quoted evidence")
    document.save(source)
    inventory = preflight_document(source, "docx").inventory
    raw = ArtifactRef(
        artifact_id="raw-pandoc",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-pandoc",
        producer_object_id="pandoc.json",
    )

    blocks, issues = _pandoc_blocks(
        {
            "blocks": [
                {"t": "Para", "c": [{"t": "Str", "c": "Opening evidence"}]},
                {
                    "t": "BlockQuote",
                    "c": [
                        {
                            "t": "Para",
                            "c": [{"t": "Str", "c": "First quoted evidence"}],
                        },
                        {
                            "t": "Para",
                            "c": [{"t": "Str", "c": "Second quoted evidence"}],
                        },
                    ],
                },
            ]
        },
        SimpleNamespace(input_snapshot_path=str(source)),
        "attempt-1",
        raw,
    )

    assert inventory["required_anchors"] == ["body/p[1]", "body/p[3]", "body/p[4]"]
    assert not issues
    assert [block.retrieval_projection for block in blocks] == [
        "Opening evidence",
        "First quoted evidence",
        "Second quoted evidence",
    ]
    assert [block.locators[0].element_path for block in blocks] == inventory["required_anchors"]

    _ = DocumentGraph(
        graph_id="graph-pandoc",
        source_sha256="a" * 64,
        input_snapshot_hash="a" * 64,
        selected_attempt_id="attempt-1",
        blocks=tuple(blocks),
        assets=(),
        issues=(),
    )



def test_pandoc_adapter_expands_numbered_docx_content_and_restores_style_headings(
    tmp_path: Path,
) -> None:
    source = tmp_path / "numbered.docx"
    document = WordDocument()
    document.add_heading("Policy", level=1)
    document.add_paragraph("Scope")
    document.add_heading("Responsibilities", level=2)
    document.add_paragraph("Review work")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Record"
    document.save(source)
    inventory = preflight_document(source, "docx").inventory
    raw = ArtifactRef(
        artifact_id="raw-pandoc",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-pandoc",
        producer_object_id="pandoc.json",
    )

    blocks, issues = _pandoc_blocks(
        {
            "blocks": [
                {
                    "t": "OrderedList",
                    "c": [
                        [1, {"t": "Decimal"}, {"t": "DefaultDelim"}],
                        [
                            [
                                {"t": "Para", "c": [{"t": "Str", "c": "Policy"}]},
                                {"t": "Para", "c": [{"t": "Str", "c": "Scope"}]},
                                {
                                    "t": "OrderedList",
                                    "c": [
                                        [1, {"t": "Decimal"}, {"t": "DefaultDelim"}],
                                        [[{"t": "Para", "c": [{"t": "Str", "c": "Responsibilities"}]}]],
                                    ],
                                },
                                {"t": "Para", "c": [{"t": "Str", "c": "Review work"}]},
                            ]
                        ],
                    ],
                },
                {"t": "Table", "c": [{"t": "Str", "c": "Record"}]},
            ]
        },
        SimpleNamespace(input_snapshot_path=str(source)),
        "attempt-1",
        raw,
    )

    assert not issues
    assert [block.kind for block in blocks] == ["heading", "list", "heading", "list", "table"]
    assert [block.payload.to_dict().get("level") for block in blocks[:3]] == [1, None, 2]
    assert [block.locators[0].element_path for block in blocks] == inventory["required_anchors"]
    _ = DocumentGraph(
        graph_id="graph-pandoc-numbered",
        source_sha256="a" * 64,
        input_snapshot_hash="a" * 64,
        selected_attempt_id="attempt-1",
        blocks=tuple(blocks),
        assets=(),
        issues=(),
    )


def test_pandoc_adapter_maps_standalone_docx_images_to_verified_obsidian_assets(
    tmp_path: Path,
) -> None:
    source = tmp_path / "figure.docx"
    document = WordDocument()
    document.add_paragraph("Figure placeholder")
    document.save(source)
    raw = ArtifactRef(
        artifact_id="raw-pandoc",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-pandoc",
        producer_object_id="pandoc.json",
    )
    image = ArtifactRef(
        artifact_id="image-artifact",
        attempt_id="attempt-1",
        sha256="b" * 64,
        media_type="image/png",
        role="image",
        private_relative_path="pending/image-artifact",
        producer_object_id="media/figure.png",
    )
    assets: list = []
    blocks, issues = _pandoc_blocks(
        {
            "blocks": [
                {
                    "t": "Para",
                    "c": [
                        {
                            "t": "Image",
                            "c": [
                                ["", [], []],
                                [{"t": "Str", "c": "Figure one"}],
                                ["media/figure.png", ""],
                            ],
                        }
                    ],
                }
            ]
        },
        SimpleNamespace(input_snapshot_path=str(source)),
        "attempt-1",
        raw,
        {image.producer_object_id: image},
        assets,
    )

    assert not issues
    assert [block.kind for block in blocks] == ["image"]
    assert len(assets) == 1
    graph = DocumentGraph(
        graph_id="graph-pandoc-image",
        source_sha256="a" * 64,
        input_snapshot_hash="a" * 64,
        selected_attempt_id="attempt-1",
        blocks=tuple(blocks),
        assets=tuple(assets),
        issues=(),
    )
    assert render_document_graph(graph).markdown == f"![[{assets[0].planned_vault_path()}|Figure one]]"


def test_pandoc_adapter_maps_absolute_windows_media_paths_to_verified_assets(
    tmp_path: Path,
) -> None:
    source = tmp_path / "figure.docx"
    document = WordDocument()
    document.add_paragraph("Figure placeholder")
    document.save(source)
    raw = ArtifactRef(
        artifact_id="raw-pandoc",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-pandoc",
        producer_object_id="pandoc.json",
    )
    image_path = tmp_path / "media" / "figure.png"
    image_path.parent.mkdir()
    image_path.write_bytes(b"png-bytes")
    image = ArtifactRef(
        artifact_id="image-artifact",
        attempt_id="attempt-1",
        sha256="b" * 64,
        media_type="image/png",
        role="image",
        private_relative_path="pending/image-artifact",
        producer_object_id="media/figure.png",
    )
    assets: list = []

    blocks, issues = _pandoc_blocks(
        {
            "blocks": [
                {
                    "t": "Para",
                    "c": [
                        {
                            "t": "Image",
                            "c": [
                                ["", [], []],
                                [{"t": "Str", "c": "Figure one"}],
                                [str(image_path), ""],
                            ],
                        }
                    ],
                }
            ]
        },
        SimpleNamespace(input_snapshot_path=str(source)),
        "attempt-1",
        raw,
        {image.producer_object_id: image},
        assets,
        artifact_root=tmp_path,
    )

    assert not issues
    assert [block.kind for block in blocks] == ["image"]
    assert len(assets) == 1


def test_mineru_content_list_adapter_uses_real_regions_and_raw_json_evidence() -> None:
    raw = ArtifactRef(
        artifact_id="raw-mineru",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-mineru",
        producer_object_id="mineru/probe_content_list_v2.json",
    )
    blocks, issues = _mineru_blocks(
        [
            [
                {
                    "type": "title",
                    "content": {"level": 2, "title_content": [{"type": "text", "content": "Title"}]},
                    "bbox": [10, 20, 100, 40],
                },
                {
                    "type": "paragraph",
                    "content": {"paragraph_content": [{"type": "text", "content": "Body"}]},
                    "bbox": [10, 50, 100, 70],
                },
            ]
        ],
        "attempt-1",
        raw,
    )

    assert not issues
    assert [block.kind for block in blocks] == ["heading", "paragraph"]
    assert isinstance(blocks[0].locators[0], PdfRegionLocator)
    assert blocks[1].retrieval_projection == "Body"
    assert blocks[0].evidence_refs[0].artifact_sha256 == raw.sha256


def test_mineru_adapter_preserves_formula_and_table_while_classifying_furniture_as_warning() -> None:
    raw = ArtifactRef(
        artifact_id="raw-mineru",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-mineru",
        producer_object_id="mineru/probe_content_list_v2.json",
    )
    blocks, issues = _mineru_blocks(
        [
            [
                {
                    "type": "equation_interline",
                    "content": {
                        "math_type": "latex",
                        "math_content": "x^2 - 1 = 0",
                    },
                    "bbox": [10, 20, 100, 40],
                },
                {
                    "type": "table",
                    "content": {
                        "html": (
                            "<table><tr><td rowspan=\"2\">Term</td><td>Meaning</td></tr>"
                            "<tr><td>Value</td></tr></table>"
                        )
                    },
                    "bbox": [10, 50, 100, 90],
                },
                {
                    "type": "page_header",
                    "content": {"page_header_content": [{"type": "text", "content": "Unit"}]},
                    "bbox": [10, 100, 100, 120],
                },
            ]
        ],
        "attempt-1",
        raw,
    )

    assert [block.kind for block in blocks] == ["formula", "table"]
    assert blocks[0].payload.to_dict()["latex"] == "x^2 - 1 = 0"
    assert blocks[1].payload.to_dict()["rows"] == [["Term", "Meaning"], ["Value"]]
    assert blocks[1].payload.to_dict()["rowspan"] == [[2, 1], [1]]
    assert len(issues) == 1
    assert issues[0].severity == "warning"

    _ = DocumentGraph(
        graph_id="graph-1",
        source_sha256="a" * 64,
        input_snapshot_hash="a" * 64,
        selected_attempt_id="attempt-1",
        blocks=tuple(blocks),
        assets=(),
        issues=tuple(issues),
    )



def test_mineru_adapter_maps_lists_and_algorithms_and_keeps_images_non_blocking() -> None:
    raw = ArtifactRef(
        artifact_id="raw-mineru",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-mineru",
        producer_object_id="mineru/probe_content_list_v2.json",
    )
    blocks, issues = _mineru_blocks(
        [[
            {
                "type": "list",
                "content": {
                    "attribute": "unordered",
                    "list_items": [
                        {"item_content": [{"type": "text", "content": "First"}]},
                        {"item_content": [{"type": "text", "content": "Second"}]},
                    ],
                },
                "bbox": [10, 20, 100, 60],
            },
            {
                "type": "index",
                "content": {
                    "list_items": [
                        {"item_content": [{"type": "text", "content": "Contents"}]},
                    ]
                },
                "bbox": [10, 70, 100, 90],
            },
            {
                "type": "algorithm",
                "content": {
                    "algorithm_content": [{"type": "text", "content": "Procedure"}],
                },
                "bbox": [10, 100, 100, 120],
            },
            {"type": "image", "content": {"image_source": {"path": "images/example.jpg"}}, "bbox": [10, 130, 100, 180]},
        ]],
        "attempt-1",
        raw,
    )

    assert [block.kind for block in blocks] == ["list", "list", "paragraph"]
    assert blocks[0].payload.to_dict()["items"] == [
        {"inline_runs": [{"kind": "text", "text": "First"}]},
        {"inline_runs": [{"kind": "text", "text": "Second"}]},
    ]
    assert [issue.code for issue in issues] == ["mineru-image-preserved-as-artifact"]
    _ = DocumentGraph(
        graph_id="graph-mineru-extended",
        source_sha256="a" * 64,
        input_snapshot_hash="a" * 64,
        selected_attempt_id="attempt-1",
        blocks=tuple(blocks),
        assets=(),
        issues=tuple(issues),
    )


def test_mineru_image_blocks_bind_verified_artifacts_and_render_obsidian_embeds() -> None:
    raw = ArtifactRef(
        artifact_id="raw-mineru",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-mineru",
        producer_object_id="mineru/auto/doc_content_list_v2.json",
    )
    image = ArtifactRef(
        artifact_id="image-artifact",
        attempt_id="attempt-1",
        sha256="b" * 64,
        media_type="image/jpeg",
        role="image",
        private_relative_path="pending/image-artifact",
        producer_object_id="mineru/auto/images/example.jpg",
    )
    assets: list = []
    blocks, issues = _mineru_blocks(
        [[
            {
                "type": "image",
                "content": {
                    "image_source": {"path": "images/example.jpg"},
                    "image_caption": [{"type": "text", "content": "Figure one"}],
                },
                "bbox": [10, 20, 100, 120],
            }
        ]],
        "attempt-1",
        raw,
        {image.producer_object_id: image},
        assets,
    )

    assert not issues
    assert [block.kind for block in blocks] == ["image"]
    assert blocks[0].retrieval_projection == "Figure one"
    assert len(assets) == 1
    assert assets[0].artifact_ref == image
    graph = DocumentGraph(
        graph_id="graph-mineru-image",
        source_sha256="a" * 64,
        input_snapshot_hash="a" * 64,
        selected_attempt_id="attempt-1",
        blocks=tuple(blocks),
        assets=tuple(assets),
        issues=(),
    )
    rendered = render_document_graph(graph)
    assert rendered.markdown == f"![[{assets[0].planned_vault_path()}|Figure one]]"


def test_artifact_collection_marks_supported_raster_outputs_as_assets(tmp_path: Path) -> None:
    image = tmp_path / "images" / "figure.jpg"
    image.parent.mkdir()
    image.write_bytes(b"jpeg-bytes")

    artifacts, drafts = _artifacts("attempt-1", tmp_path)

    assert len(artifacts) == len(drafts) == 1
    assert artifacts[0].role == "image"
    assert artifacts[0].media_type == "image/jpeg"


def test_mineru_graph_adapter_connects_collected_image_artifacts(tmp_path: Path) -> None:
    output = tmp_path / "mineru" / "auto" / "doc_content_list_v2.json"
    image = tmp_path / "mineru" / "auto" / "images" / "figure.jpg"
    image.parent.mkdir(parents=True)
    output.write_text(
        json.dumps(
            [[
                {
                    "type": "image",
                    "content": {"image_source": {"path": "images/figure.jpg"}},
                    "bbox": [10, 20, 100, 120],
                }
            ]]
        ),
        encoding="utf-8",
    )
    image.write_bytes(b"jpeg-bytes")
    artifacts, _ = _artifacts("attempt-1", tmp_path)
    raw = next(
        artifact
        for artifact in artifacts
        if artifact.producer_object_id == "mineru/auto/doc_content_list_v2.json"
    )

    graph = _adapt_graph(
        "mineru",
        output,
        SimpleNamespace(source_sha256="a" * 64, input_snapshot_hash="a" * 64),
        "attempt-1",
        raw,
        artifacts,
    )

    assert [block.kind for block in graph.blocks] == ["image"]
    assert len(graph.assets) == 1
    assert graph.assets[0].artifact_ref.role == "image"


def test_mineru_empty_paragraph_becomes_a_located_review_issue() -> None:
    raw = ArtifactRef(
        artifact_id="raw-mineru",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-mineru",
        producer_object_id="mineru/probe_content_list_v2.json",
    )

    blocks, issues = _mineru_blocks(
        [[{"type": "paragraph", "content": {"paragraph_content": []}, "bbox": [10, 20, 100, 40]}]],
        "attempt-1",
        raw,
    )

    assert not blocks
    assert [issue.code for issue in issues] == ["mineru-empty-text"]
    assert isinstance(issues[0].locator, PdfRegionLocator)


def test_docling_empty_formula_becomes_a_located_review_issue_instead_of_raising() -> None:
    raw = ArtifactRef(
        artifact_id="raw-docling",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-docling",
        producer_object_id="docling/result.json",
    )
    blocks, issues = _docling_blocks(
        {
            "texts": [
                {
                    "self_ref": "#/texts/0",
                    "label": "formula",
                    "orig": "x^2 - 1 = 0",
                    "text": "",
                    "prov": [
                        {
                            "page_no": 1,
                            "bbox": {"l": 10, "b": 20, "r": 100, "t": 40},
                        }
                    ],
                }
            ]
        },
        "pdf",
        "attempt-1",
        raw,
    )

    assert not blocks
    assert [issue.code for issue in issues] == ["docling-formula-unresolved"]
    assert isinstance(issues[0].locator, PdfRegionLocator)


def test_docling_adapter_omits_repeated_margin_furniture_and_derives_numbered_heading_levels() -> None:
    raw = ArtifactRef(
        artifact_id="raw-docling",
        attempt_id="attempt-1",
        sha256="a" * 64,
        media_type="application/json",
        role="converter-json",
        private_relative_path="pending/raw-docling",
        producer_object_id="docling/result.json",
    )

    def text(text: str, label: str, page: int, bottom: float, top: float) -> dict[str, object]:
        return {
            "self_ref": f"#/texts/{text}-{page}",
            "label": label,
            "text": text,
            "prov": [{"page_no": page, "bbox": {"l": 10, "b": bottom, "r": 100, "t": top}}],
        }

    blocks, issues = _docling_blocks(
        {
            "pages": {
                "1": {"size": {"height": 100}},
                "2": {"size": {"height": 100}},
                "3": {"size": {"height": 100}},
            },
            "texts": [
                text("Controlled procedure", "section_header", 1, 91, 98),
                text("Controlled procedure", "section_header", 2, 91, 98),
                text("Controlled procedure", "section_header", 3, 91, 98),
                text("Page 1 of 3", "text", 1, 2, 8),
                text("Page 2 of 3", "text", 2, 2, 8),
                text("Page 3 of 3", "text", 3, 2, 8),
                text("4 Duties", "section_header", 1, 50, 58),
                text("Work must be reviewed.", "text", 1, 38, 45),
                text("4.1 Supervisor", "section_header", 2, 50, 58),
                text("4.1.1 Check records", "section_header", 3, 50, 58),
                text("continued.", "section_header", 3, 38, 45),
            ],
        },
        "pdf",
        "attempt-1",
        raw,
    )

    assert [block.retrieval_projection for block in blocks] == [
        "4 Duties",
        "Work must be reviewed.",
        "4.1 Supervisor",
        "4.1.1 Check records",
        "continued.",
    ]
    assert [block.kind for block in blocks] == ["heading", "paragraph", "heading", "heading", "paragraph"]
    assert [block.payload.to_dict().get("level") for block in blocks if block.kind == "heading"] == [
        1,
        2,
        3,
    ]
    assert [issue.code for issue in issues] == ["docling-page-furniture-omitted"] * 6
    markdown = render_document_graph(
        DocumentGraph(
            graph_id="graph-docling-furniture",
            source_sha256="a" * 64,
            input_snapshot_hash="a" * 64,
            selected_attempt_id="attempt-1",
            blocks=tuple(blocks),
            assets=(),
            issues=tuple(issues),
        )
    ).markdown
    assert "Controlled procedure" not in markdown
    assert "# 4 Duties" in markdown
    assert "## 4.1 Supervisor" in markdown
    assert "### 4.1.1 Check records" in markdown


def _write_manifest(root: Path, profile: dict[str, object]) -> None:
    (root / "converter-profiles.json").write_text(
        json.dumps({"schema_version": 1, "profiles": [profile]}), encoding="utf-8"
    )


def _digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()
